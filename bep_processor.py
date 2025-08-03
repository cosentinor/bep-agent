"""
Main BEP processing module that orchestrates the entire workflow.
"""

import os
import json
from typing import List, Dict, Tuple, Optional
from document_loader import DocumentLoader
from vector_store import VectorStore
from outline_generator import OutlineGenerator
from config import Config

class BEPProcessor:
    """Main class that orchestrates the BEP generation workflow."""
    
    def __init__(self):
        """Initialize the BEP processor with all components."""
        self.document_loader = DocumentLoader()
        self.vector_store = VectorStore()
        self.outline_generator = OutlineGenerator()
        
        self.documents = []
        self.chunks = []
        self.is_initialized = False
    
    def initialize(self, force_rebuild: bool = False) -> bool:
        """
        Initialize the BEP processor by loading documents and building vector store.
        
        Args:
            force_rebuild: Whether to force rebuilding the vector store
            
        Returns:
            True if initialization successful, False otherwise
        """
        try:
            # Validate configuration
            Config.validate_config()
            
            # Try to load existing vector store if not forcing rebuild
            if not force_rebuild and self.vector_store.load_vector_store():
                print("Loaded existing vector store")
                self.is_initialized = True
                return True
            
            # Load BEP documents
            print("Loading BEP documents...")
            self.documents = self.document_loader.load_bep_documents()
            
            if not self.documents:
                print("No BEP documents found. Please add .docx files to the data/bep_samples directory.")
                return False
            
            # Create text chunks
            print("Creating text chunks...")
            self.chunks = self.document_loader.create_text_chunks(self.documents)
            
            if not self.chunks:
                print("No text chunks created from documents.")
                return False
            
            # Build vector store
            print("Building vector store...")
            self.vector_store.build_faiss_index(self.chunks)
            self.vector_store.compute_plan_vectors(self.documents)
            
            # Save vector store for future use
            self.vector_store.save_vector_store()
            
            self.is_initialized = True
            print("BEP processor initialized successfully!")
            return True
            
        except Exception as e:
            print(f"Error initializing BEP processor: {str(e)}")
            return False
    
    def get_document_summaries(self) -> List[Dict]:
        """
        Get summaries of loaded BEP documents.
        
        Returns:
            List of document summary dictionaries
        """
        summaries = []
        for doc in self.documents:
            summaries.append({
                'filename': doc['filename'],
                'word_count': doc['word_count'],
                'heading_count': len(doc['headings']),
                'sample_headings': doc['headings'][:5]  # First 5 headings
            })
        return summaries
    
    def analyze_project_requirements(self, project_description: str,
                                   requirements_files: List[str] = None,
                                   additional_bep_files: List[str] = None) -> Dict:
        """
        Analyze project requirements and find similar BEPs.

        Args:
            project_description: Text description of the project
            requirements_files: Optional list of paths to requirements documents
            additional_bep_files: Optional list of additional BEP files to include

        Returns:
            Analysis results dictionary
        """
        if not self.is_initialized:
            raise RuntimeError("BEP processor not initialized. Call initialize() first.")

        # Process additional BEP files if provided
        additional_documents = []
        if additional_bep_files:
            for bep_file in additional_bep_files:
                if os.path.exists(bep_file):
                    print(f"Processing additional BEP: {os.path.basename(bep_file)}")
                    # Load the additional BEP document
                    if bep_file.endswith('.docx'):
                        text, headings = self.document_loader.extract_text_from_docx(bep_file)
                    elif bep_file.endswith('.pdf'):
                        text, headings = self.document_loader.extract_text_from_pdf(bep_file)
                    else:
                        continue

                    if text:
                        additional_documents.append({
                            'filename': os.path.basename(bep_file),
                            'file_path': bep_file,
                            'text': text,
                            'headings': headings,
                            'word_count': len(text.split()),
                            'file_type': 'docx' if bep_file.endswith('.docx') else 'pdf',
                            'is_temporary': True
                        })

        # Load additional requirements if files provided
        full_requirements = project_description
        requirements_content = []

        if requirements_files:
            for req_file in requirements_files:
                if os.path.exists(req_file):
                    file_content = self.document_loader.load_project_requirements(
                        file_path=req_file
                    )
                    if file_content:
                        requirements_content.append(f"[From {os.path.basename(req_file)}]\n{file_content}")

        if requirements_content:
            full_requirements = f"{project_description}\n\n" + "\n\n".join(requirements_content)
        
        # Find similar plans
        similar_plans = self.vector_store.find_similar_plans(full_requirements)
        
        # Find relevant chunks
        relevant_chunks = self.vector_store.search_similar_chunks(full_requirements)
        
        return {
            'project_description': project_description,
            'full_requirements': full_requirements,
            'similar_plans': similar_plans,
            'relevant_chunks': relevant_chunks,
            'requirements_files': requirements_files or [],
            'additional_bep_files': additional_bep_files or [],
            'additional_documents': additional_documents,
            'requirements_count': len(requirements_files) if requirements_files else 0,
            'additional_bep_count': len(additional_documents)
        }
    
    def generate_bep_outline(self, analysis_results: Dict, 
                           selected_plans: List[str] = None) -> List[str]:
        """
        Generate a BEP outline based on analysis results.
        
        Args:
            analysis_results: Results from analyze_project_requirements
            selected_plans: Optional list of specific plans to use as reference
            
        Returns:
            List of outline headings
        """
        if not self.is_initialized:
            raise RuntimeError("BEP processor not initialized. Call initialize() first.")
        
        # Use selected plans or top similar plans
        if not selected_plans:
            selected_plans = [plan[0] for plan in analysis_results['similar_plans'][:2]]
        
        # Extract canonical outline from selected plans
        canonical_headings = self.outline_generator.extract_canonical_outline(
            self.documents, selected_plans
        )
        
        # Generate project-specific outline
        outline = self.outline_generator.generate_project_outline(
            analysis_results['full_requirements'],
            analysis_results['relevant_chunks'],
            canonical_headings
        )
        
        return outline
    
    def generate_bep_content(self, analysis_results: Dict, outline: List[str]) -> Dict[str, str]:
        """
        Generate complete BEP content for all sections.
        
        Args:
            analysis_results: Results from analyze_project_requirements
            outline: List of section headings
            
        Returns:
            Dictionary mapping section headings to content
        """
        if not self.is_initialized:
            raise RuntimeError("BEP processor not initialized. Call initialize() first.")
        
        # Generate content for all sections
        bep_content = self.outline_generator.generate_complete_bep(
            analysis_results['full_requirements'],
            outline,
            analysis_results['relevant_chunks']
        )
        
        return bep_content
    
    def export_bep_markdown(self, outline: List[str], content: Dict[str, str], 
                           project_name: str = "New Project") -> str:
        """
        Export BEP as markdown format.
        
        Args:
            outline: List of section headings
            content: Dictionary mapping headings to content
            project_name: Name of the project
            
        Returns:
            Markdown formatted BEP
        """
        markdown_lines = [
            f"# BIM Execution Plan - {project_name}",
            "",
            f"*Generated by BEP Agent*",
            "",
            "## Table of Contents",
            ""
        ]
        
        # Add table of contents
        for i, heading in enumerate(outline, 1):
            markdown_lines.append(f"{i}. [{heading}](#{heading.lower().replace(' ', '-')})")
        
        markdown_lines.extend(["", "---", ""])
        
        # Add sections
        for i, heading in enumerate(outline, 1):
            markdown_lines.extend([
                f"## {i}. {heading}",
                "",
                content.get(heading, "Content to be developed."),
                "",
                "---",
                ""
            ])
        
        return "\n".join(markdown_lines)
    
    def export_bep_json(self, outline: List[str], content: Dict[str, str], 
                       analysis_results: Dict, project_name: str = "New Project") -> str:
        """
        Export BEP as JSON format.
        
        Args:
            outline: List of section headings
            content: Dictionary mapping headings to content
            analysis_results: Analysis results for metadata
            project_name: Name of the project
            
        Returns:
            JSON formatted BEP
        """
        bep_data = {
            "project_name": project_name,
            "generated_by": "BEP Agent",
            "metadata": {
                "similar_plans": analysis_results.get('similar_plans', []),
                "relevant_chunks_count": len(analysis_results.get('relevant_chunks', [])),
                "project_description": analysis_results.get('project_description', '')
            },
            "outline": outline,
            "sections": []
        }
        
        for i, heading in enumerate(outline, 1):
            section = {
                "number": i,
                "heading": heading,
                "content": content.get(heading, "Content to be developed.")
            }
            bep_data["sections"].append(section)
        
        return json.dumps(bep_data, indent=2)

    def export_bep_docx(self, outline: List[str], content: Dict[str, str],
                       analysis_results: Dict, project_name: str = "New Project",
                       template_path: str = None, output_path: str = None) -> str:
        """
        Export BEP as DOCX format using template.

        Args:
            outline: List of section headings
            content: Dictionary mapping headings to content
            analysis_results: Analysis results for metadata
            project_name: Name of the project
            template_path: Path to DOCX template (optional)
            output_path: Path for output DOCX file (optional)

        Returns:
            Path to created DOCX file
        """
        try:
            from docx_writer import DOCXWriter
            import tempfile
            import json
            from datetime import datetime

            # Create JSON data for template replacement
            bep_data = {
                "Project_Name": project_name,
                "Generated_Date": datetime.now().strftime("%Y-%m-%d"),
                "Generated_By": "BEP Agent - Atkinsrealis",
                "Version": "1.3.0"
            }

            # Add section content
            for i, heading in enumerate(outline, 1):
                section_key = f"Section_{i}_Title"
                content_key = f"Section_{i}_Content"
                bep_data[section_key] = heading
                bep_data[content_key] = content.get(heading, "Content to be developed.")

            # Add analysis metadata
            if analysis_results.get('similar_plans'):
                similar_plans = [plan[0] for plan in analysis_results['similar_plans'][:3]]
                bep_data["Similar_Plans"] = ", ".join(similar_plans)

            # Create temporary JSON file
            with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False) as temp_json:
                json.dump(bep_data, temp_json, indent=2)
                temp_json_path = temp_json.name

            # Set output path if not provided
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = f"BEP_{project_name.replace(' ', '_')}_{timestamp}.docx"

            # Use template if provided, otherwise create basic document
            if template_path and os.path.exists(template_path):
                writer = DOCXWriter()
                writer.process_document(template_path, temp_json_path, output_path)
            else:
                # Create basic DOCX if no template provided
                self._create_basic_docx(bep_data, output_path)

            # Clean up temporary file
            os.unlink(temp_json_path)

            return output_path

        except Exception as e:
            print(f"Error exporting DOCX: {str(e)}")
            return None

    def _create_basic_docx(self, data: Dict, output_path: str) -> None:
        """Create a basic DOCX document without template."""
        from docx import Document
        from docx.shared import Inches

        doc = Document()

        # Add title
        title = doc.add_heading(f"BIM Execution Plan - {data.get('Project_Name', 'New Project')}", 0)

        # Add metadata
        doc.add_paragraph(f"Generated: {data.get('Generated_Date', '')}")
        doc.add_paragraph(f"Generated by: {data.get('Generated_By', '')}")
        doc.add_paragraph("")

        # Add sections
        section_num = 1
        while f"Section_{section_num}_Title" in data:
            title_key = f"Section_{section_num}_Title"
            content_key = f"Section_{section_num}_Content"

            # Add section heading
            doc.add_heading(f"{section_num}. {data[title_key]}", level=1)

            # Add section content
            doc.add_paragraph(data[content_key])
            doc.add_paragraph("")  # Add spacing

            section_num += 1

        doc.save(output_path)
    
    def get_processing_stats(self) -> Dict:
        """
        Get statistics about the processing state.
        
        Returns:
            Dictionary with processing statistics
        """
        return {
            "is_initialized": self.is_initialized,
            "documents_loaded": len(self.documents),
            "chunks_created": len(self.chunks),
            "vector_store_size": self.vector_store.faiss_index.ntotal if self.vector_store.faiss_index else 0,
            "plan_vectors_computed": len(self.vector_store.plan_vectors),
            "embedding_model": "Local" if Config.USE_LOCAL_EMBEDDINGS else "OpenAI"
        }
