"""
DOCX Writer Module for BEP Agent.
Converts JSON output into fully formatted Word documents using BEP templates.
"""

import json
import argparse
import sys
from pathlib import Path
from typing import Dict, Any, Set
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.text.run import Run
from docx.table import Table


class DOCXWriter:
    """Handles conversion of JSON data to formatted DOCX using templates."""
    
    def __init__(self):
        """Initialize the DOCX writer."""
        self.replacement_keys: Set[str] = set()
        self.replacements: Dict[str, str] = {}
    
    def load_json_data(self, json_path: str) -> Dict[str, Any]:
        """
        Load and flatten JSON data for template replacement.
        
        Args:
            json_path: Path to JSON file
            
        Returns:
            Flattened dictionary with placeholder keys
        """
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Flatten the JSON structure to match placeholder format
            flattened = self._flatten_json(data)
            
            # Convert to placeholder format (add curly braces if not present)
            self.replacements = {}
            for key, value in flattened.items():
                # Ensure key is in placeholder format {{key}}
                if not key.startswith('{{') or not key.endswith('}}'):
                    placeholder_key = f"{{{{{key}}}}}"
                else:
                    placeholder_key = key
                
                self.replacements[placeholder_key] = str(value) if value is not None else ""
            
            self.replacement_keys = set(self.replacements.keys())
            print(f"Loaded {len(self.replacements)} replacements from JSON")
            return self.replacements
            
        except Exception as e:
            raise ValueError(f"Error loading JSON file {json_path}: {str(e)}")
    
    def _flatten_json(self, data: Any, parent_key: str = '', separator: str = '_') -> Dict[str, Any]:
        """
        Flatten nested JSON structure.
        
        Args:
            data: JSON data to flatten
            parent_key: Parent key for nested items
            separator: Separator for nested keys
            
        Returns:
            Flattened dictionary
        """
        items = []
        
        if isinstance(data, dict):
            for key, value in data.items():
                new_key = f"{parent_key}{separator}{key}" if parent_key else key
                
                if isinstance(value, (dict, list)):
                    items.extend(self._flatten_json(value, new_key, separator).items())
                else:
                    items.append((new_key, value))
        
        elif isinstance(data, list):
            for i, value in enumerate(data):
                new_key = f"{parent_key}{separator}{i}" if parent_key else str(i)
                
                if isinstance(value, (dict, list)):
                    items.extend(self._flatten_json(value, new_key, separator).items())
                else:
                    items.append((new_key, value))
        
        else:
            items.append((parent_key, data))
        
        return dict(items)
    
    def process_document(self, template_path: str, json_path: str, output_path: str) -> None:
        """
        Process template document with JSON data replacements.
        
        Args:
            template_path: Path to template DOCX file
            json_path: Path to JSON data file
            output_path: Path for output DOCX file
        """
        try:
            # Load JSON data
            self.load_json_data(json_path)
            
            # Load template document
            print(f"Loading template: {template_path}")
            doc = Document(template_path)
            
            # Process paragraphs
            self._process_paragraphs(doc)
            
            # Process tables
            self._process_tables(doc)
            
            # Save output document
            print(f"Saving output: {output_path}")
            doc.save(output_path)
            print(f"Successfully created {output_path}")
            
        except Exception as e:
            raise RuntimeError(f"Error processing document: {str(e)}")
    
    def _process_paragraphs(self, doc: Document) -> None:
        """Process all paragraphs in the document for replacements."""
        for paragraph in doc.paragraphs:
            self._process_runs(paragraph.runs)
    
    def _process_tables(self, doc: Document) -> None:
        """Process all tables in the document for replacements."""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_runs(paragraph.runs)
    
    def _process_runs(self, runs: list) -> None:
        """
        Process runs for placeholder replacement and highlight removal.
        
        Args:
            runs: List of runs to process
        """
        i = 0
        while i < len(runs):
            run = runs[i]
            
            # Check if run has yellow highlight or contains placeholder
            is_highlighted = self._is_yellow_highlighted(run)
            contains_placeholder = any(key in run.text for key in self.replacement_keys)
            
            if is_highlighted or contains_placeholder:
                # Process this run for replacements
                new_text = run.text
                
                # Replace all matching placeholders
                for placeholder, replacement in self.replacements.items():
                    if placeholder in new_text:
                        new_text = new_text.replace(placeholder, replacement)
                        print(f"Replaced '{placeholder}' with '{replacement[:50]}...'")
                
                # Update run text and clear highlight
                run.text = new_text
                if is_highlighted:
                    run.font.highlight_color = None
            
            i += 1
    
    def _is_yellow_highlighted(self, run: Run) -> bool:
        """
        Check if a run has yellow highlighting.
        
        Args:
            run: Run to check
            
        Returns:
            True if run has yellow highlight
        """
        try:
            return run.font.highlight_color == WD_COLOR_INDEX.YELLOW
        except Exception:
            return False


def cli():
    """Command line interface for DOCX Writer."""
    parser = argparse.ArgumentParser(
        description="Convert JSON data to formatted DOCX using templates"
    )
    parser.add_argument(
        '--template', 
        required=True, 
        help='Path to template DOCX file with placeholders'
    )
    parser.add_argument(
        '--json', 
        required=True, 
        help='Path to JSON file with replacement data'
    )
    parser.add_argument(
        '--out', 
        required=True, 
        help='Path for output DOCX file'
    )
    
    args = parser.parse_args()
    
    # Validate input files exist
    if not Path(args.template).exists():
        print(f"Error: Template file not found: {args.template}")
        sys.exit(1)
    
    if not Path(args.json).exists():
        print(f"Error: JSON file not found: {args.json}")
        sys.exit(1)
    
    # Create output directory if needed
    output_path = Path(args.out)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    try:
        # Process document
        writer = DOCXWriter()
        writer.process_document(args.template, args.json, args.out)
        print("Document processing completed successfully!")
        
    except Exception as e:
        print(f"Error: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    cli()
