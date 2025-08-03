"""
Structure extractor module for document outline learning.
Extracts hierarchical structure from DOCX and PDF documents.
"""

import os
import re
from typing import List, Dict, Optional, Union
from dataclasses import dataclass
from docx import Document
from docx.shared import Inches
import fitz  # PyMuPDF


@dataclass
class Node:
    """Represents a document structure node (heading/section)."""
    id: str
    level: int
    title: str
    text_span: str
    doc_id: str


class StructureExtractor:
    """Extracts hierarchical structure from documents."""
    
    def __init__(self):
        """Initialize the structure extractor."""
        self.supported_formats = {'.docx', '.pdf'}
    
    def extract(self, doc_path: str) -> List[Node]:
        """
        Auto-detect document type and extract structure.
        
        Args:
            doc_path: Path to the document
            
        Returns:
            List of Node objects representing document structure
            
        Raises:
            ValueError: If document format is not supported
            FileNotFoundError: If document file doesn't exist
        """
        if not os.path.exists(doc_path):
            raise FileNotFoundError(f"Document not found: {doc_path}")
        
        file_ext = os.path.splitext(doc_path)[1].lower()
        
        if file_ext == '.docx':
            return self.parse_docx(doc_path)
        elif file_ext == '.pdf':
            return self.parse_pdf(doc_path)
        else:
            raise ValueError(f"Unsupported document format: {file_ext}. "
                           f"Supported formats: {', '.join(self.supported_formats)}")
    
    def parse_docx(self, path: str) -> List[Node]:
        """
        Extract structure from DOCX document.
        
        Args:
            path: Path to DOCX file
            
        Returns:
            List of Node objects representing document structure
        """
        try:
            doc = Document(path)
            nodes = []
            doc_id = os.path.basename(path)
            
            # Extract headings using paragraph styles
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.style.name.startswith('Heading'):
                    # Extract heading level from style name
                    level = self._extract_heading_level(paragraph.style.name)
                    title = paragraph.text.strip()
                    
                    if title:  # Only add non-empty headings
                        # Get text span (content following this heading)
                        text_span = self._get_text_span_docx(doc, i)
                        
                        node = Node(
                            id=f"{doc_id}_h{i}_{level}",
                            level=level,
                            title=title,
                            text_span=text_span,
                            doc_id=doc_id
                        )
                        nodes.append(node)
            
            # If no styled headings found, try regex-based extraction
            if not nodes:
                nodes = self._extract_headings_regex_docx(doc, doc_id)
            
            return nodes
            
        except Exception as e:
            print(f"Error parsing DOCX file {path}: {str(e)}")
            return []
    
    def parse_pdf(self, path: str) -> List[Node]:
        """
        Extract structure from PDF document.
        
        Args:
            path: Path to PDF file
            
        Returns:
            List of Node objects representing document structure
        """
        try:
            doc = fitz.open(path)
            nodes = []
            doc_id = os.path.basename(path)
            
            # Try to extract from table of contents first
            toc = doc.get_toc()
            if toc:
                nodes = self._extract_from_toc(toc, doc, doc_id)
            
            # If no TOC, try text-based extraction
            if not nodes:
                nodes = self._extract_headings_regex_pdf(doc, doc_id)
            
            doc.close()
            return nodes
            
        except Exception as e:
            print(f"Error parsing PDF file {path}: {str(e)}")
            return []
    
    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        # Extract number from style name like "Heading 1", "Heading 2", etc.
        match = re.search(r'(\d+)', style_name)
        return int(match.group(1)) if match else 1
    
    def _get_text_span_docx(self, doc: Document, heading_index: int, max_chars: int = 500) -> str:
        """
        Get text content following a heading in DOCX document.
        
        Args:
            doc: Document object
            heading_index: Index of the heading paragraph
            max_chars: Maximum characters to extract
            
        Returns:
            Text content following the heading
        """
        text_parts = []
        char_count = 0
        
        # Start from the paragraph after the heading
        for i in range(heading_index + 1, len(doc.paragraphs)):
            paragraph = doc.paragraphs[i]
            
            # Stop if we hit another heading
            if paragraph.style.name.startswith('Heading'):
                break
            
            text = paragraph.text.strip()
            if text:
                if char_count + len(text) > max_chars:
                    # Truncate to fit within max_chars
                    remaining = max_chars - char_count
                    text = text[:remaining] + "..."
                    text_parts.append(text)
                    break
                
                text_parts.append(text)
                char_count += len(text)
        
        return ' '.join(text_parts)
    
    def _extract_headings_regex_docx(self, doc: Document, doc_id: str) -> List[Node]:
        """Extract headings using regex patterns from DOCX."""
        nodes = []
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Common heading patterns
        patterns = [
            (r'^(\d+\.?\s+.+)$', 1),  # "1. Title" or "1 Title"
            (r'^(\d+\.\d+\.?\s+.+)$', 2),  # "1.1. Title" or "1.1 Title"
            (r'^(\d+\.\d+\.\d+\.?\s+.+)$', 3),  # "1.1.1. Title"
            (r'^([A-Z][A-Z\s]{2,})$', 1),  # ALL CAPS headings
        ]
        
        lines = full_text.split('\n')
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            for pattern, level in patterns:
                if re.match(pattern, line, re.MULTILINE):
                    # Get surrounding text as span
                    text_span = self._get_text_span_from_lines(lines, i)
                    
                    node = Node(
                        id=f"{doc_id}_regex_{i}_{level}",
                        level=level,
                        title=line,
                        text_span=text_span,
                        doc_id=doc_id
                    )
                    nodes.append(node)
                    break
        
        return nodes
    
    def _extract_from_toc(self, toc: List, doc, doc_id: str) -> List[Node]:
        """Extract structure from PDF table of contents."""
        nodes = []
        
        for i, (level, title, page_num) in enumerate(toc):
            # Get text from the page
            try:
                page = doc[page_num - 1]  # PyMuPDF uses 0-based indexing
                page_text = page.get_text()
                
                # Extract a portion of text as span
                text_span = page_text[:500] + "..." if len(page_text) > 500 else page_text
                
                node = Node(
                    id=f"{doc_id}_toc_{i}_{level}",
                    level=level,
                    title=title.strip(),
                    text_span=text_span,
                    doc_id=doc_id
                )
                nodes.append(node)
                
            except Exception as e:
                print(f"Error extracting text for TOC entry {title}: {str(e)}")
                continue
        
        return nodes
    
    def _extract_headings_regex_pdf(self, doc, doc_id: str) -> List[Node]:
        """Extract headings using regex patterns from PDF."""
        nodes = []
        full_text = ""
        
        # Extract all text from PDF
        for page in doc:
            full_text += page.get_text() + "\n"
        
        # Use same patterns as DOCX
        patterns = [
            (r'^(\d+\.?\s+.+)$', 1),
            (r'^(\d+\.\d+\.?\s+.+)$', 2),
            (r'^(\d+\.\d+\.\d+\.?\s+.+)$', 3),
            (r'^([A-Z][A-Z\s]{2,})$', 1),
        ]
        
        lines = full_text.split('\n')
        for i, line in enumerate(lines):
            line = line.strip()
            if not line or len(line) < 3:
                continue
                
            for pattern, level in patterns:
                if re.match(pattern, line, re.MULTILINE):
                    text_span = self._get_text_span_from_lines(lines, i)
                    
                    node = Node(
                        id=f"{doc_id}_pdf_regex_{i}_{level}",
                        level=level,
                        title=line,
                        text_span=text_span,
                        doc_id=doc_id
                    )
                    nodes.append(node)
                    break
        
        return nodes
    
    def _get_text_span_from_lines(self, lines: List[str], heading_index: int, max_chars: int = 500) -> str:
        """Get text span following a heading from list of lines."""
        text_parts = []
        char_count = 0
        
        for i in range(heading_index + 1, len(lines)):
            line = lines[i].strip()
            if not line:
                continue
            
            # Stop if we hit another potential heading
            if re.match(r'^(\d+\.?\s+.+)$', line) or re.match(r'^([A-Z][A-Z\s]{2,})$', line):
                break
            
            if char_count + len(line) > max_chars:
                remaining = max_chars - char_count
                line = line[:remaining] + "..."
                text_parts.append(line)
                break
            
            text_parts.append(line)
            char_count += len(line)
        
        return ' '.join(text_parts)


# Convenience functions for direct use
def parse_docx(path: str) -> List[Node]:
    """Parse DOCX document and return structure nodes."""
    extractor = StructureExtractor()
    return extractor.parse_docx(path)


def parse_pdf(path: str) -> List[Node]:
    """Parse PDF document and return structure nodes."""
    extractor = StructureExtractor()
    return extractor.parse_pdf(path)


def extract(doc_path: str) -> List[Node]:
    """Auto-detect document type and extract structure."""
    extractor = StructureExtractor()
    return extractor.extract(doc_path)
