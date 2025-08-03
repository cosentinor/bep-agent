"""
Unit tests for DOCX Writer module.
"""

import json
import tempfile
import unittest
from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx_writer import DOCXWriter


class TestDOCXWriter(unittest.TestCase):
    """Test cases for DOCX Writer functionality."""
    
    def setUp(self):
        """Set up test fixtures."""
        self.writer = DOCXWriter()
        self.temp_dir = Path(tempfile.mkdtemp())
    
    def tearDown(self):
        """Clean up test fixtures."""
        # Clean up temporary files
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def create_test_template(self) -> Path:
        """
        Create a test template with placeholders and yellow highlights.
        
        Returns:
            Path to created template file
        """
        doc = Document()
        
        # Add paragraph with placeholder
        p1 = doc.add_paragraph()
        run1 = p1.add_run("Project Name: {{Project_Name}}")
        run1.font.highlight_color = WD_COLOR_INDEX.YELLOW
        
        # Add another paragraph with different placeholder
        p2 = doc.add_paragraph()
        run2 = p2.add_run("Location: {{Project_Location}}")
        run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
        
        # Add normal text (should not be changed)
        p3 = doc.add_paragraph("This is normal text that should not change.")
        
        # Add table with placeholder
        table = doc.add_table(rows=2, cols=2)
        cell = table.cell(0, 0)
        cell_run = cell.paragraphs[0].add_run("Budget: {{Project_Budget}}")
        cell_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        
        # Save template
        template_path = self.temp_dir / "test_template.docx"
        doc.save(str(template_path))
        return template_path
    
    def create_test_json(self) -> Path:
        """
        Create test JSON data file.
        
        Returns:
            Path to created JSON file
        """
        test_data = {
            "Project_Name": "Waterfront Redevelopment – Phase 2",
            "Project_Location": "Downtown Marina District",
            "Project_Budget": "$2.5 Million",
            "unused_field": "This should not appear"
        }
        
        json_path = self.temp_dir / "test_data.json"
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(test_data, f, indent=2)
        
        return json_path
    
    def test_load_json_data(self):
        """Test JSON data loading and flattening."""
        json_path = self.create_test_json()
        
        replacements = self.writer.load_json_data(str(json_path))
        
        # Check that data was loaded and formatted correctly
        self.assertIn("{{Project_Name}}", replacements)
        self.assertEqual(replacements["{{Project_Name}}"], "Waterfront Redevelopment – Phase 2")
        self.assertIn("{{Project_Location}}", replacements)
        self.assertEqual(replacements["{{Project_Location}}"], "Downtown Marina District")
    
    def test_flatten_json(self):
        """Test JSON flattening functionality."""
        nested_data = {
            "project": {
                "name": "Test Project",
                "details": {
                    "budget": 1000000,
                    "timeline": "12 months"
                }
            },
            "team": ["Alice", "Bob", "Charlie"]
        }
        
        flattened = self.writer._flatten_json(nested_data)
        
        # Check flattened structure
        self.assertIn("project_name", flattened)
        self.assertEqual(flattened["project_name"], "Test Project")
        self.assertIn("project_details_budget", flattened)
        self.assertEqual(flattened["project_details_budget"], 1000000)
        self.assertIn("team_0", flattened)
        self.assertEqual(flattened["team_0"], "Alice")
    
    def test_process_document(self):
        """Test complete document processing."""
        # Create test files
        template_path = self.create_test_template()
        json_path = self.create_test_json()
        output_path = self.temp_dir / "output.docx"
        
        # Process document
        self.writer.process_document(str(template_path), str(json_path), str(output_path))
        
        # Verify output file was created
        self.assertTrue(output_path.exists())
        
        # Load and inspect output document
        output_doc = Document(str(output_path))
        
        # Check that placeholders were replaced
        all_text = []
        for paragraph in output_doc.paragraphs:
            all_text.append(paragraph.text)
        
        # Check table content
        for table in output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        all_text.append(paragraph.text)
        
        full_text = " ".join(all_text)
        
        # Verify replacements occurred
        self.assertIn("Waterfront Redevelopment – Phase 2", full_text)
        self.assertIn("Downtown Marina District", full_text)
        self.assertIn("$2.5 Million", full_text)
        
        # Verify placeholders were removed
        self.assertNotIn("{{Project_Name}}", full_text)
        self.assertNotIn("{{Project_Location}}", full_text)
        self.assertNotIn("{{Project_Budget}}", full_text)
        
        # Verify normal text was preserved
        self.assertIn("This is normal text that should not change", full_text)
    
    def test_highlight_detection(self):
        """Test yellow highlight detection."""
        doc = Document()
        paragraph = doc.add_paragraph()
        
        # Create run with yellow highlight
        highlighted_run = paragraph.add_run("Highlighted text")
        highlighted_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        
        # Create normal run
        normal_run = paragraph.add_run("Normal text")
        
        # Test detection
        self.assertTrue(self.writer._is_yellow_highlighted(highlighted_run))
        self.assertFalse(self.writer._is_yellow_highlighted(normal_run))
    
    def test_placeholder_replacement(self):
        """Test placeholder replacement in runs."""
        # Set up test replacements
        self.writer.replacements = {
            "{{test_placeholder}}": "Replacement Text",
            "{{another_placeholder}}": "Another Replacement"
        }
        self.writer.replacement_keys = set(self.writer.replacements.keys())
        
        # Create document with placeholders
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("Text with {{test_placeholder}} and {{another_placeholder}}")
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        
        # Process runs
        self.writer._process_runs(paragraph.runs)
        
        # Check replacement
        self.assertEqual(
            paragraph.runs[0].text, 
            "Text with Replacement Text and Another Replacement"
        )
        
        # Check highlight was removed
        self.assertIsNone(paragraph.runs[0].font.highlight_color)


if __name__ == "__main__":
    unittest.main()
